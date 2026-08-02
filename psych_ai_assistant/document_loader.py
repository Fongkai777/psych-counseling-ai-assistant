from io import BytesIO
import re
import subprocess
import tempfile
from pathlib import Path


def extract_uploaded_text(uploaded_file):
    name = uploaded_file.name
    suffix = Path(name).suffix.lower()
    data = uploaded_file.getvalue()

    if suffix == ".pdf":
        return name, extract_pdf(data)
    if suffix == ".docx":
        return name, extract_docx(data)
    if suffix == ".doc":
        return name, extract_doc(data, name)
    if suffix in {".txt", ".md"}:
        return name, data.decode("utf-8", errors="ignore")

    raise ValueError(f"暂不支持 {suffix or '未知格式'} 文件。")


def extract_pdf(data):
    plumber_text = extract_pdf_with_pdfplumber(data)
    pypdf_text = extract_pdf_with_pypdf(data)
    candidates = [text for text in (plumber_text, pypdf_text) if text.strip()]
    page_count = pdf_page_count(data)
    if not candidates:
        if page_count:
            raise RuntimeError(
                f"这个 PDF 共 {page_count} 页，但没有可提取的文字层，"
                "大概率是扫描版/图片版 PDF，需要先做 OCR 再导入。"
            )
        raise RuntimeError("这个 PDF 没有可提取的文字内容，需要先做 OCR 再导入。")
    text = max(candidates, key=len)
    if page_count >= 20 and len(text) < page_count * 80:
        raise RuntimeError(
            f"这个 PDF 共 {page_count} 页，但只提取到 {len(text)} 字，"
            "正文大概率是扫描图片，当前只读到了少量元信息，需要先做 OCR 再导入。"
        )
    return text


def pdf_page_count(data):
    try:
        from pypdf import PdfReader
    except ImportError:
        return 0
    try:
        return len(PdfReader(BytesIO(data)).pages)
    except Exception:
        return 0


def extract_pdf_with_pdfplumber(data):
    try:
        import pdfplumber
    except ImportError:
        return ""

    pages = []
    try:
        with pdfplumber.open(BytesIO(data)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(layout=True) or page.extract_text() or ""
                text = normalize_extracted_text(text)
                if text.strip():
                    pages.append(f"第 {index} 页\n{text}")
    except Exception:
        return ""
    return "\n\n".join(pages).strip()


def extract_pdf_with_pypdf(data):
    try:
        from pypdf import PdfReader
        from pypdf.errors import DependencyError
    except ImportError as exc:
        raise RuntimeError("缺少 pypdf，请先安装依赖。") from exc

    try:
        reader = PdfReader(BytesIO(data))
    except DependencyError as exc:
        raise RuntimeError(
            "这个 PDF 使用了 AES 加密，需要 cryptography 依赖。"
            "如果刚安装过依赖，请重启 Streamlit 后再上传。"
        ) from exc
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except DependencyError as exc:
            raise RuntimeError(
                "这个 PDF 的页面内容需要 cryptography 解密。"
                "如果刚安装过依赖，请重启 Streamlit 后再上传。"
            ) from exc
        text = normalize_extracted_text(text)
        if text.strip():
            pages.append(f"第 {index} 页\n{text}")
    return "\n\n".join(pages).strip()


def normalize_extracted_text(text):
    lines = []
    cleaned = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", "", text or "")
    for line in cleaned.replace("\r", "\n").splitlines():
        line = " ".join(line.split())
        if looks_like_noise(line):
            continue
        if line:
            lines.append(line)
    return "\n".join(lines).strip()


def looks_like_noise(line):
    if re.search(r"QQ\s*\d{5,}", line, flags=re.I):
        return True
    if len(line) < 4:
        return False
    chinese = len(re.findall(r"[\u4e00-\u9fff]", line))
    ascii_alnum = len(re.findall(r"[A-Za-z0-9]", line))
    useful = chinese + ascii_alnum
    return useful / max(len(line), 1) < 0.35


def extract_docx(data):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，请先安装依赖。") from exc

    document = Document(BytesIO(data))
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def extract_doc(data, name):
    with tempfile.TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / name
        path.write_bytes(data)
        try:
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", str(path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=30,
            )
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise RuntimeError("这个 .doc 文件无法自动解析，建议另存为 .docx 或 PDF 后再导入。") from exc
    return result.stdout.strip()
