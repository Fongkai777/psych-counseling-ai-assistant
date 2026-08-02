from __future__ import annotations

import argparse
import re
import statistics
from collections import Counter
from pathlib import Path

from docx import Document


RELATION_WORDS = (
    "同事",
    "朋友",
    "老师",
    "领导",
    "室友",
    "同学",
    "前任",
    "对象",
    "家人",
    "爸爸",
    "妈妈",
    "父亲",
    "母亲",
)

STYLE_MARKERS = (
    "我觉得",
    "我感觉",
    "其实",
    "可能",
    "也许",
    "但是",
    "所以",
    "而且",
    "某种程度",
    "说到底",
    "说实话",
    "认真说",
    "有点",
    "反而",
    "大概",
    "只是",
    "还是",
    "确实",
    "本质上",
    "后来",
    "突然",
)


def extract_docx(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def anonymize(text: str) -> str:
    text = re.sub(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", "[邮箱]", text, flags=re.I)
    text = re.sub(r"(?<!\d)1[3-9]\d{9}(?!\d)", "[手机号]", text)
    text = re.sub(r"\b[A-Za-z][A-Za-z0-9_.-]{2,20}\b", "[英文名/账号]", text)
    text = re.sub(r"微信[:：]?\s*[\w.-]{4,}", "微信：[账号]", text, flags=re.I)
    text = re.sub(r"QQ[:：]?\s*\d{5,12}", "QQ：[账号]", text, flags=re.I)
    for word in RELATION_WORDS:
        text = re.sub(rf"({word})([叫是为])?[\u4e00-\u9fff]{{2,4}}", rf"\1[某人]", text)
    text = re.sub(r"小[\u4e00-\u9fff]", "小某", text)
    text = re.sub(r"老[\u4e00-\u9fff]", "老某", text)
    return text


def split_sentences(text: str) -> list[str]:
    pieces = re.split(r"(?<=[。！？!?；;])\s*", text)
    return [piece.strip() for piece in pieces if piece.strip()]


def marker_summary(text: str) -> list[tuple[str, int]]:
    counts = Counter()
    for marker in STYLE_MARKERS:
        count = text.count(marker)
        if count:
            counts[marker] = count
    return counts.most_common(12)


def build_profile(source_dir: Path) -> str:
    files = sorted(source_dir.glob("*.docx"))
    records = []
    raw_texts = []
    clean_texts = []
    for path in files:
        raw = extract_docx(path)
        clean = anonymize(raw)
        raw_texts.append(raw)
        clean_texts.append(clean)
        records.append((path.name, len(raw), len(split_sentences(raw))))

    corpus = "\n\n".join(clean_texts)
    sentences = split_sentences(corpus)
    sentence_lengths = [len(sentence) for sentence in sentences if sentence]
    avg_len = statistics.mean(sentence_lengths) if sentence_lengths else 0
    median_len = statistics.median(sentence_lengths) if sentence_lengths else 0
    markers = marker_summary(corpus)
    first_person_count = sum(corpus.count(word) for word in ("我", "自己", "我的"))
    question_count = corpus.count("？") + corpus.count("?")
    exclamation_count = corpus.count("！") + corpus.count("!")

    marker_text = "\n".join(f"- `{word}`：{count} 次" for word, count in markers) or "- 暂无明显高频口头标记"
    file_rows = "\n".join(
        f"- {name}：约 {char_count} 字，{sentence_count} 个句子"
        for name, char_count, sentence_count in records
    )

    return f"""# 个人语气画像（已脱敏）

> 来源：`{source_dir}` 下的 Word 日记文件。本文档只保留可用于写作风格迁移的抽象特征，不保留原始日记全文；人名、账号、联系方式等已做规则脱敏。

## 资料概览

- 文件数：{len(files)}
- 脱敏后总字数：约 {len(corpus)} 字
- 句子数：{len(sentences)}
- 平均句长：{avg_len:.1f} 字
- 句长中位数：{median_len:.1f} 字
- 第一人称密度参考：`我 / 自己 / 我的` 共出现 {first_person_count} 次
- 问号：{question_count} 次；感叹号：{exclamation_count} 次

## 文件清单

{file_rows}

## 高频语气标记

{marker_text}

## 可迁移的表达习惯

- 允许保留第一人称视角，像是在认真复盘一件事，而不是站在高处给别人下判断。
- 多使用“我觉得 / 可能 / 其实 / 有点 / 某种程度上”这类带有试探和自我校准的表达。
- 适合先承认复杂性，再给出判断；不要一上来给结论。
- 可以把抽象问题拆成“我当时怎么感受到它”“它为什么让我卡住”“我后来怎么理解它”。
- 语气可以有一点自嘲和诚实，但不要油滑，不要鸡汤。
- 适合用短句承接情绪，再用稍长的句子做解释。
- 不要把话说满。多使用“可能是”“更像是”“我会把它理解成”“不一定是”。
- 适合把建议写成很小、很具体、可以马上做的动作，而不是宏大的自我改造口号。

## 回答时的模仿原则

1. 用“我会先把这个问题理解为……”这类开场，把问题放进一种个人化的理解框架。
2. 不要假装绝对客观；可以承认“这只是我的理解”。
3. 少用整齐但冷的模板句，多用自然转折：其实、但是、所以、后来、某种程度上。
4. 每段只解决一个小意思，段落不要太长。
5. 在表达共情时保持克制，不要过度安慰，不要说“你一定会好起来”。
6. 结尾可以轻一点，但不要硬升华；更像是“给自己留一个可执行的入口”。

## 不建议模仿的部分

- 不模仿具体经历、具体关系和私人事实。
- 不复述日记里的私人事件。
- 不使用未脱敏的人名、地点、账号或联系方式。
- 不为了像本人而输出过度低落、攻击性或容易引发误读的表达。
"""


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    profile = build_profile(args.source_dir)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(profile, encoding="utf-8")
    print(args.out)


if __name__ == "__main__":
    main()
